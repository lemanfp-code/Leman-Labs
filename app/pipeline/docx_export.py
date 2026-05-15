"""Export .docx — reproduction de LA FORME des dossiers CPC / CRY.

Charte tirée des PDF de référence (references/ des repos) :
- Couverture plein fond sombre (CPC noir + orange ; CRY navy + blanc italique)
- Bandeau d'en-tête couleur programme « DOSSIER <MOIS ANNÉE> »
- Titres de bloc en bandeau pleine largeur
- Filet vertical à gauche, corps justifié, pagination en pastille
La FORME = structure des blocs (prompt) + cette identité visuelle.
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from programs import Program

BODY_FONT = "Calibri"
TITLE_FONT = "Calibri"
DARK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

ILLUS_RE = re.compile(r"【\s*Illustration.*?】")
TOKEN_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*)")

# Palette par programme (issue des PDF de référence)
STYLES = {
    "cpc": {
        "accent": "E8792B", "cover_bg": "0F0F1A",
        "cover_title": "F59E4F", "cover_text": "E6E6EC",
        "brand": "CLUB PRIVÉ CRYPTO", "brand_bg": "E8792B", "brand_fg": "FFFFFF",
        "header_bg": "E8792B", "band_bg": "FBE2CC", "band_fg": "C75F12",
        "title_italic": False, "kicker": None,
    },
    "cry": {
        "accent": "0E7490", "cover_bg": "0A1428",
        "cover_title": "FFFFFF", "cover_text": "9FB6D8",
        "brand": "CRYPTOS EXPONENTIELLES", "brand_bg": None, "brand_fg": "7DD3FC",
        "header_bg": "15294D", "band_bg": "1F3A5F", "band_fg": "FFFFFF",
        "title_italic": True, "kicker": "— SYNTHÈSE EXCLUSIVE —",
    },
    "_default": {
        "accent": "1F3A93", "cover_bg": "12233F",
        "cover_title": "FFFFFF", "cover_text": "AFC0DA",
        "brand": "", "brand_bg": None, "brand_fg": "FFFFFF",
        "header_bg": "1F3A93", "band_bg": "1F3A93", "band_fg": "FFFFFF",
        "title_italic": False, "kicker": None,
    },
}


def _style_for(program: Program) -> dict:
    st = dict(STYLES.get(program.id, STYLES["_default"]))
    if program.docx_accent:
        st["accent"] = program.docx_accent.lstrip("#").upper()
    if not st.get("brand"):
        st["brand"] = program.name.upper()
    return st


def _shade(parent_elem, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    parent_elem.append(shd)


def _cell_shade(cell, hex_color: str):
    _shade(cell._tc.get_or_add_tcPr(), hex_color)


def _para_border(paragraph, edges: dict):
    """edges: {'left':('20','E8792B'), 'bottom':(...)} → (size, color)."""
    pPr = paragraph._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    for side, (sz, color) in edges.items():
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), sz)
        e.set(qn("w:space"), "8" if side in ("left", "right") else "2")
        e.set(qn("w:color"), color)
        bdr.append(e)
    pPr.append(bdr)


def _content_width(section) -> int:
    return section.page_width - section.left_margin - section.right_margin


def _full_table(doc, section):
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    t.allow_autofit = False
    w = _content_width(section)
    t.columns[0].width = w
    t.rows[0].cells[0].width = w
    tblPr = t._tbl.tblPr
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(w))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    return t


def _run(p, text, *, size, color, bold=False, italic=False, font=BODY_FONT):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = color if isinstance(color, RGBColor) else RGBColor.from_string(color)
    r.font.bold = bold
    r.font.italic = italic
    return r


def _add_runs(paragraph, text, color=DARK, base_size=10.5, base_bold=False):
    first = True
    for line in text.split("\n"):
        if not first:
            paragraph.add_run().add_break()
        first = False
        for tok in TOKEN_RE.split(line):
            if not tok:
                continue
            bold, ital, inner = base_bold, False, tok
            if tok.startswith("**") and tok.endswith("**"):
                bold, inner = True, tok[2:-2]
            elif tok.startswith("*") and tok.endswith("*"):
                ital, inner = True, tok[1:-1]
            _run(paragraph, inner, size=base_size, color=color, bold=bold, italic=ital)


def _page_field(paragraph, st):
    run = paragraph.add_run()
    run.font.name = TITLE_FONT
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = WHITE
    _shade(run._r.get_or_add_rPr(), st["accent"])
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _cover(doc, section, st, program, month, year, punchline):
    t = _full_table(doc, section)
    cell = t.rows[0].cells[0]
    _cell_shade(cell, st["cover_bg"])
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "center")
    tcPr.append(va)
    trPr = t.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement("w:trHeight")
    th.set(qn("w:val"), str(int(Cm(23.5))))
    th.set(qn("w:hRule"), "atLeast")
    trPr.append(th)

    cell.text = ""
    if st["kicker"]:
        k = cell.paragraphs[0]
        k.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(k, st["kicker"], size=12, color=st["cover_text"], bold=True, font=TITLE_FONT)
    else:
        tag = cell.paragraphs[0]
        tag.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(tag, f"  DOSSIER {month.upper()} {year}  ", size=11,
             color=st["brand_fg"], bold=True, font=TITLE_FONT)
        if st["brand_bg"]:
            _shade(tag._p.get_or_add_pPr(), st["brand_bg"])

    cell.add_paragraph()
    tp = cell.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(10)
    tp.paragraph_format.space_after = Pt(10)
    _run(tp, (punchline or f"DOSSIER {month.upper()} {year}").upper(),
         size=30, color=st["cover_title"], bold=True,
         italic=st["title_italic"], font=TITLE_FONT)

    dp = cell.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(dp, f"  {month.upper()} {year}  ", size=13, color=st["cover_text"],
         bold=True, font=TITLE_FONT)

    cell.add_paragraph()
    bp = cell.add_paragraph()
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.paragraph_format.space_before = Pt(28)
    _run(bp, f"  {st['brand']}  ", size=15, color=st["brand_fg"],
         bold=True, font=TITLE_FONT)
    if st["brand_bg"]:
        _shade(bp._p.get_or_add_pPr(), st["brand_bg"])

    doc.add_section(WD_SECTION.NEW_PAGE)


def _header_footer(section, st, month, year):
    # Section de contenu indépendante de la couverture
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    # En-tête : bandeau couleur programme + DOSSIER MOIS ANNÉE (blanc, à droite)
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _shade(hp._p.get_or_add_pPr(), st["header_bg"])
    _run(hp, f"DOSSIER {month.upper()} {year}   ", size=10,
         color=WHITE, bold=True, font=TITLE_FONT)
    # Pied : numéro de page en pastille accent, à droite
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _page_field(fp, st)


def _band(doc, section, st, text):
    text = text.strip()
    if text.startswith("**") and text.endswith("**"):
        text = text[2:-2]
    t = _full_table(doc, section)
    cell = t.rows[0].cells[0]
    _cell_shade(cell, st["band_bg"])
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    _run(p, text.upper(), size=15, color=st["band_fg"], bold=True,
         italic=st["title_italic"], font=TITLE_FONT)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)


def _subhead(doc, st, text, level):
    text = text.strip()
    if text.startswith("**") and text.endswith("**"):
        text = text[2:-2]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 3 else 8)
    p.paragraph_format.space_after = Pt(4)
    _run(p, text, size=12.5 if level == 3 else 11.5,
         color=RGBColor.from_string(st["accent"]), bold=True,
         italic=st["title_italic"], font=TITLE_FONT)


def _illustration(doc, st, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    _shade(p._p.get_or_add_pPr(), "F2F2F2")
    _para_border(p, {"left": ("18", st["accent"])})
    _run(p, "📷  " + text.strip(), size=10, color=RGBColor.from_string(st["accent"]),
         bold=True, italic=True)


def _table(doc, header, rows, st):
    tbl = doc.add_table(rows=1, cols=len(header))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        c = tbl.rows[0].cells[j]
        c.text = ""
        _cell_shade(c, st["accent"])
        _run(c.paragraphs[0], h.strip(), size=9.5, color=WHITE, bold=True)
    for row in rows:
        cells = tbl.add_row().cells
        for j in range(len(header)):
            cells[j].text = ""
            _add_runs(cells[j].paragraphs[0], row[j] if j < len(row) else "",
                      color=DARK, base_size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _extract_punchline(lines):
    """1er titre Markdown = accroche éditoriale (mise en couverture, retirée du corps)."""
    for idx, l in enumerate(lines):
        m = re.match(r"^#{1,2}\s+(.*)$", l.strip())
        if m:
            t = m.group(1).strip().strip("*")
            if not re.match(r"(?i)^dossier\b", t):
                return t, idx
            return None, idx
    return None, -1


def build_docx(markdown_text: str, program: Program, month: str, year: str, out_path: str) -> str:
    st = _style_for(program)
    accent_rgb = RGBColor.from_string(st["accent"])

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK

    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(1.4)
    sec.left_margin = sec.right_margin = Cm(1.6)

    lines = (markdown_text or "").replace("\r", "").split("\n")
    punchline, skip_idx = _extract_punchline(lines)

    _cover(doc, sec, st, program, month, year, punchline)

    body = doc.sections[-1]
    body.top_margin = Cm(2.2)
    body.bottom_margin = Cm(1.8)
    body.left_margin = body.right_margin = Cm(2.0)
    _header_footer(body, st, month, year)

    i = 0
    while i < len(lines):
        if i == skip_idx:
            i += 1
            continue
        l = lines[i]

        if re.match(r"^\s*\|.*\|\s*$", l) and i + 1 < len(lines) and re.match(r"^\s*\|?[\s:|-]+\|[\s:|-]*$", lines[i + 1]):
            header = [c.strip() for c in l.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and re.match(r"^\s*\|.*\|\s*$", lines[i]):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            _table(doc, header, rows, st)
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", l)
        if m:
            lvl = len(m.group(1))
            if lvl <= 2:
                _band(doc, body, st, m.group(2))
            else:
                _subhead(doc, st, m.group(2), min(lvl, 4))
            i += 1
            continue

        if ILLUS_RE.search(l):
            _illustration(doc, st, ILLUS_RE.search(l).group(0))
            i += 1
            continue

        if re.match(r"^\s*[-*]\s+", l):
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                p = doc.add_paragraph(style="List Bullet")
                _add_runs(p, re.sub(r"^\s*[-*]\s+", "", lines[i]))
                i += 1
            continue

        if re.match(r"^\s*\d+\.\s+", l):
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                p = doc.add_paragraph(style="List Number")
                _add_runs(p, re.sub(r"^\s*\d+\.\s+", "", lines[i]))
                i += 1
            continue

        if l.strip() == "" or re.match(r"^\s*---\s*$", l):
            i += 1
            continue

        para = []
        while i < len(lines) and lines[i].strip() != "" and i != skip_idx and not re.match(r"^(#{1,6}\s|\s*[-*]\s|\s*\d+\.\s|\s*\|)", lines[i]) and not ILLUS_RE.search(lines[i]):
            para.append(lines[i])
            i += 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.12
        _add_runs(p, "\n".join(para))

    doc.core_properties.title = f"Dossier {program.short} {month} {year}"
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path
